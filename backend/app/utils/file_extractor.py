"""
从上传的文件中提取纯文本，供风格分析使用。
支持：PDF / DOCX / TXT / MD / 图片（OpenAI Vision OCR）。
"""
import base64
import io
import logging
from typing import Optional

from app.core.config import settings

logger = logging.getLogger(__name__)

TEXT_EXTS = {".txt", ".md", ".markdown"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
DOC_EXTS = {".doc"}  # 老版 Word，单独处理给出明确提示
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

SUPPORTED_EXTS = TEXT_EXTS | PDF_EXTS | DOCX_EXTS | IMAGE_EXTS


class UnsupportedFileType(Exception):
    pass


def _ext(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as e:
            logger.warning(f"PDF 解密失败（可能加密）: {e}")

    pages = []
    for idx, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            pages.append(text)
            logger.info(f"PDF 第 {idx+1} 页提取字符数: {len(text)}")
        except Exception as e:
            logger.warning(f"PDF 第 {idx+1} 页提取失败: {e}")
    result = "\n".join(p for p in pages if p).strip()
    logger.info(f"PDF 共 {len(reader.pages)} 页，合计提取 {len(result)} 字符")
    return result


async def _ocr_pdf_via_vision(data: bytes) -> str:
    """扫描版 PDF 兜底：逐页转图片走 Vision OCR。需要 pdf2image + poppler。"""
    if not settings.OPENAI_API_KEY:
        logger.warning("未配置 OPENAI_API_KEY，跳过 PDF OCR 兜底")
        return ""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.warning("未安装 pdf2image，无法对扫描 PDF 做 OCR。pip install pdf2image 并安装 poppler")
        return ""

    try:
        images = convert_from_bytes(data, dpi=150)
    except Exception as e:
        logger.error(f"PDF 转图片失败: {e}")
        return ""

    pieces = []
    for idx, img in enumerate(images[:20]):  # 限 20 页，防超额
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        text = await _extract_image_via_vision(buf.getvalue(), "image/png")
        logger.info(f"PDF OCR 第 {idx+1} 页字符数: {len(text)}")
        if text:
            pieces.append(text)
    return "\n".join(pieces).strip()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        logger.error(f"python-docx 未安装: {e}")
        raise

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        logger.error(f"DOCX 打开失败（可能不是合法 .docx，或文件被加密）: {e}")
        raise

    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    result = "\n".join(parts).strip()
    logger.info(f"DOCX 提取段落 {len(parts)} 段，合计 {len(result)} 字符")
    return result


def _extract_text(data: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


async def _extract_image_via_vision(data: bytes, mime: str) -> str:
    """用 OpenAI Vision 对图片 OCR / 内容描述。"""
    if not settings.OPENAI_API_KEY:
        logger.warning("未配置 OPENAI_API_KEY，跳过图片 OCR")
        return ""

    try:
        import openai

        client = openai.AsyncOpenAI(
            api_key=settings.OPENAI_API_KEY,
            base_url=settings.OPENAI_API_BASE,
        )
        b64 = base64.b64encode(data).decode()
        data_url = f"data:{mime or 'image/png'};base64,{b64}"

        # 选择支持视觉的模型；DEFAULT_AI_MODEL 可能不支持，使用 gpt-4o-mini 兜底
        model = settings.DEFAULT_AI_MODEL
        if "gpt-4" not in model and "gpt-4o" not in model:
            model = "gpt-4o-mini"

        resp = await client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "请提取图片中的全部文字内容，按原文排版输出，不要添加任何解释。如果图片不含文字，请简要描述图片所表达的主题与风格特征。",
                        },
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            temperature=0,
            max_tokens=1500,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logger.error(f"图片 OCR 失败: {e}")
        return ""


async def extract_text(
    *,
    filename: str,
    data: bytes,
    content_type: Optional[str] = None,
) -> str:
    """
    根据文件名后缀分发到对应解析器。返回提取到的纯文本。
    """
    ext = _ext(filename)
    if ext in TEXT_EXTS:
        return _extract_text(data)
    if ext in PDF_EXTS:
        text = _extract_pdf(data)
        if not text or len(text) < 20:
            logger.warning(f"PDF 文本提取结果过短({len(text)})，尝试 OCR 兜底")
            ocr_text = await _ocr_pdf_via_vision(data)
            if ocr_text:
                return ocr_text
        return text
    if ext in DOCX_EXTS:
        return _extract_docx(data)
    if ext in IMAGE_EXTS:
        return await _extract_image_via_vision(data, content_type or "image/png")
    if ext in DOC_EXTS:
        raise UnsupportedFileType(
            "不支持老版 .doc 格式，请用 Word/WPS 另存为 .docx 后再上传"
        )
    raise UnsupportedFileType(f"不支持的文件类型: {ext or filename}")
